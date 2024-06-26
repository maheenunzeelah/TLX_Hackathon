from fastapi import FastAPI, File, UploadFile,Form
import pickle
from fastapi.middleware.cors import CORSMiddleware
from roboflow import display_detected_objects
import uvicorn
from PIL import Image
import requests
import tempfile
import io
import os
from fastapi.responses import JSONResponse
import base64
from app import generateResult,generatePromptResult
from docx import Document
from embeddings import generateEmbeddings
from dalle import generate_image
app=FastAPI(debug=True)
origins = [
   
    "http://localhost:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"message": "Hello World"}
@app.post('/detectedElements')    
async def generateBoundingBoxes(upload_file: UploadFile = File(...)):
   
    contents = await upload_file.read()
    image = Image.open(io.BytesIO(contents))

    # Convert the image to a format suitable for Roboflow
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
    detected_elements , annotated_img= display_detected_objects(img_str)
    return {'annotated_img':annotated_img,'detected_elements':detected_elements}

@app.post('/apiExtract')    
async def generateApis(upload_file: UploadFile = File(...),userStory: UploadFile = File(...)):
    try:
        # Read the file content into memory
        user_story = await userStory.read()

        # Load the content into a Document object using python-docx
        doc = Document(io.BytesIO(user_story))

        # Extract text from the document
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Join all paragraphs into a single string
        user_story = "\n".join(text)
        print(user_story)
        embb=generateEmbeddings(user_story)
        print(embb)
       
        user_query="Given the detected elements, given image and user story, extract possible apis with their methods, sample endpints, response and request bodies. The answer should be properly structured and in json"
        prompt=f"Context: {embb} \nUser: {user_query}\nAssistant:"
       
        contents = await upload_file.read()
        image = Image.open(io.BytesIO(contents))
        print(prompt,"prompt")
        # Convert the image to a format suitable for Roboflow
        if image.mode == 'RGBA':
            image = image.convert('RGB')
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
    
        # print(detected_elements,"detected_elements")
        detected_elements , _= display_detected_objects(img_str)
        result=generateResult(detected_elements,prompt,img_str)
        print(result,"resultt")
    # Return the inferenceresults as a JSON response
        return {'result':result}


        # Return the extracted text as JSON response
    #     return JSONResponse(content={"filename": userStory.filename, "content": full_text})
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    

@app.post('/samplePrompts')    
async def generatePrompt(prompt:str=Form(...)):
   

    result=generatePromptResult(prompt)

    # Return the inferenceresults as a JSON response
    return {'result':result}

@app.post('/code')    
async def generateCode(upload_file: UploadFile = File(...),userPrompt:str=Form(...)):
    print(userPrompt)
    prompt="Generate HTML and inline CSS code for a signup page in English. Include fields for username, email, password, and a signup button. Ensure the page design is clean and user-friendly, with appropriate styling for input fields and a responsive layout."
    if userPrompt =='no':
        prompt="Given the detected elements and given image, write it's html and inline css code."
    
    contents = await upload_file.read()
    image = Image.open(io.BytesIO(contents))

    # Convert the image to a format suitable for Roboflow
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
 
    # print(detected_elements,"detected_elements")
    detected_elements , _= display_detected_objects(img_str)
    result=generateResult(detected_elements,prompt,img_str)

    # Return the inferenceresults as a JSON response
    return {'result':result}
    
@app.post('/suggestions')    
async def generateSuggestion(upload_file: UploadFile = File(...),userStory: UploadFile = File(...)):
    try:
        # Read the file content into memory
        user_story = await userStory.read()

        # Load the content into a Document object using python-docx
        doc = Document(io.BytesIO(user_story))

        # Extract text from the document
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Join all paragraphs into a single string
        user_story = "\n".join(text)
        print(user_story)
        embb=generateEmbeddings(user_story)
        print(embb)
       
        user_query="Given the detected elements, given image and user story, generate improvement suggestions"
        prompt=f"Context: {embb} \nUser: {user_query}\nAssistant:"
       
        contents = await upload_file.read()
        image = Image.open(io.BytesIO(contents))
        print(prompt,"prompt")
        # Convert the image to a format suitable for Roboflow
        if image.mode == 'RGBA':
            image = image.convert('RGB')
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
    
        # print(detected_elements,"detected_elements")
        detected_elements , _= display_detected_objects(img_str)
        result=generateResult(detected_elements,prompt,img_str)
        print(result,"resultt")
    # Return the inferenceresults as a JSON response
        return {'result':result}


        # Return the extracted text as JSON response
    #     return JSONResponse(content={"filename": userStory.filename, "content": full_text})
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    
@app.post("/generateImage")
async def generateImages(userStory: UploadFile = File(...),user_prompt:str=Form(...)):
    try:
        # Read the file content into memory
        user_story = await userStory.read()

        # Load the content into a Document object using python-docx
        doc = Document(io.BytesIO(user_story))

        # Extract text from the document
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Join all paragraphs into a single string
        user_story = "\n".join(text)
        # print(user_story)
        # embb=generateEmbeddings(user_story)
        # print(embb)
       
        user_query="Given the user story, and the following prompt, generate image" +user_prompt
        prompt=f" {user_story}\n  {user_query}"
    
        result=generate_image(user_query)
        print(result,"resultt")
    # Return the inferenceresults as a JSON response
        return {'result':result}


        # Return the extracted text as JSON response
    #     return JSONResponse(content={"filename": userStory.filename, "content": full_text})
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    
if __name__ == '__main__':
    uvicorn.run(app, host='127.0.0.1', port=4000)   